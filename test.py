import pickle
import os
import cv2
import streamlit as st
import mediapipe as mp
import pandas as pd
import numpy as np
from datetime import datetime
import docx
import wave
import pyaudio
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import whisper
import PyPDF2

# Initialize the necessary objects and models
mp_drawing = mp.solutions.drawing_utils
mp_holistic = mp.solutions.holistic

# Load the body language detection model
with open('model.pkl', 'rb') as file:
    model = pickle.load(file)

holistic = mp_holistic.Holistic(min_detection_confidence=0.5, min_tracking_confidence=0.5)

# Load pre-trained Sentence-BERT model for job description and transcript matching
embedding_model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Load Whisper model for speech-to-text transcription
whisper_model = whisper.load_model("base")

# Job descriptions for different positions
job_descriptions = {
    "Data Scientist": "We are looking for a data scientist with experience in machine learning, Python, and data analysis. The ideal candidate should be able to build predictive models and analyze complex datasets.",
    "Software Engineer": "We are looking for a software engineer with expertise in programming languages like Python, Java, and C++. Experience with system design and large-scale distributed systems is required.",
    "Product Manager": "The ideal product manager should have experience in agile development, stakeholder management, and product lifecycle management. Familiarity with market research and data analysis is a plus.",
    "Data Analyst": "We are looking for a data analyst with expertise in data visualization, SQL, and statistical analysis. The candidate should be able to turn data into actionable insights and reports."
}

# Set up Streamlit
st.title("Interview Emotion and Pose Detection with Webcam")

# Sidebar for interviewee and resume selection
st.sidebar.title("Select Interviewee and Resume")
interviewees = ["Hoang", "Fidan", "Ojashwi"]
resume_paths = {
    "Hoang": "C:/Users/dang0/Downloads/hoangResume.pdf",
    "Fidan": "fidan.pdf",
    "Ojashwi": "ojashwi.pdf"
}
selected_interviewee = st.sidebar.selectbox("Choose an Interviewee:", interviewees)

# Upload resume button
uploaded_resume = st.sidebar.file_uploader("Upload Resume", type=['pdf'])

# Declare session state for managing references
if 'pdf_ref' not in st.session_state:
    st.session_state.pdf_ref = None
if uploaded_resume:
    st.session_state.pdf_ref = uploaded_resume  # backup

# Display resume if uploaded
if st.session_state.pdf_ref:
    binary_data = st.session_state.pdf_ref.getvalue()
    st.write(binary_data)

# Layout: live webcam on the left, result and transcript on the right and bottom
col1, col2 = st.columns(2)
col3 = st.container()

with col1:
    FRAME_WINDOW = st.image([])

with col2:
    predictions_container = st.empty()

with col3:
    st.write("### Interview Transcript")
    transcript_container = st.empty()

# Declare variable for predictions
predictions = []

# Initialize flags for recording and saving
is_recording = False
transcript = []
pose_data = []

# Function to capture audio from the microphone
def capture_audio():
    p = pyaudio.PyAudio()
    stream = p.open(format=pyaudio.paInt16,
                    channels=1,
                    rate=16000,
                    input=True,
                    frames_per_buffer=1024)

    frames = []
    for i in range(0, int(16000 / 1024 * 5)):  # Record for 5 seconds
        data = stream.read(1024)
        frames.append(data)
    
    stream.stop_stream()
    stream.close()
    p.terminate()

    # Save the recorded audio to a WAV file
    audio_path = "audio_recording.wav"
    wf = wave.open(audio_path, 'wb')
    wf.setnchannels(1)
    wf.setsampwidth(p.get_sample_size(pyaudio.paInt16))
    wf.setframerate(16000)
    wf.writeframes(b''.join(frames))
    wf.close()

    return audio_path

# Function to transcribe audio using Whisper
def transcribe_audio(audio_path):
    result = whisper_model.transcribe(audio_path)
    return result['text']

# Function to extract text from the uploaded PDF resume
def extract_resume_text(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Save transcript to a Word document
def save_transcript_to_file(transcript_text):
    doc = docx.Document()
    doc.add_heading("Interview Transcript", 0)

    for time, text in transcript_text:
        doc.add_paragraph(f"{time} - {text}")

    doc_path = f"transcript_{selected_interviewee}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(doc_path)
    st.write(f"Transcript saved to {doc_path}")

# Save pose detection data to Excel
def save_pose_to_excel(pose_data):
    df = pd.DataFrame(pose_data)
    excel_path = f"pose_detection_{selected_interviewee}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    df.to_excel(excel_path, index=False)
    st.write(f"Pose detection results saved to {excel_path}")

# Function to start recording
def start_recording():
    global is_recording, pose_data, transcript
    is_recording = True
    pose_data = []  # Reset pose data
    transcript = []  # Reset transcript

    cap = cv2.VideoCapture(0)  # Open webcam

    while is_recording:
        ret, frame = cap.read()
        if not ret:
            st.write("Failed to grab frame.")
            break

        # Recolor the frame (BGR to RGB)
        image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        image.flags.writeable = False

        # Make detections
        results = holistic.process(image)

        image.flags.writeable = True
        image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)

        # Draw landmarks for face, hands, and pose
        mp_drawing.draw_landmarks(image, results.face_landmarks, mp_holistic.FACEMESH_TESSELATION)
        mp_drawing.draw_landmarks(image, results.pose_landmarks, mp_holistic.POSE_CONNECTIONS)

        # Extract pose and face landmarks
        try:
            pose = results.pose_landmarks.landmark
            pose_row = list(np.array([[landmark.x, landmark.y, landmark.z, landmark.visibility] for landmark in pose]).flatten())

            # Prediction for body language
            X = pd.DataFrame([pose_row])
            body_language_class = model.predict(X)[0]
            body_language_prob = model.predict_proba(X)[0]

            # Store results
            timestamp = datetime.now().strftime("%H:%M:%S")
            pose_data.append({"Time": timestamp, "Class": body_language_class, "Probability": round(body_language_prob[np.argmax(body_language_prob)], 2)})

            # Capture and transcribe audio
            audio_path = capture_audio()
            transcript_text = transcribe_audio(audio_path)

            # Store transcript and body language class with timestamp
            transcript.append((timestamp, body_language_class, transcript_text))

        except:
            pass

        # Update webcam feed in Streamlit
        FRAME_WINDOW.image(image)

        # Display the detection results
        with predictions_container:
            st.write("### Detection Results")
            st.write(pd.DataFrame(pose_data))

        # Display the transcript
        with transcript_container:
            st.write("### Interview Transcript")
            for time, class_text, transcribed_text in transcript:
                st.write(f"{time} - {class_text}: {transcribed_text}")

    cap.release()  # Release the webcam when done

# Function to calculate similarity between text
def calculate_similarity(resume_text, transcript_text):
    resume_embedding = embedding_model.encode([resume_text])
    transcript_embedding = embedding_model.encode([transcript_text])

    similarity = cosine_similarity(resume_embedding, transcript_embedding)
    return similarity[0][0]

# Control buttons for starting and stopping the recording
if st.button("Start Recording"):
    start_recording()

if st.button("Stop Recording"):
    is_recording = False
    save_transcript_to_file(transcript)
    save_pose_to_excel(pose_data)

# Analyze transcript button
if st.button("Analyze Transcript"):
    if uploaded_resume:
        resume_text = extract_resume_text(st.session_state.pdf_ref)
        full_transcript = "\n".join([t[2] for t in transcript])
        similarity_score = calculate_similarity(resume_text, full_transcript)
        st.write(f"Similarity Score between Resume and Transcript: {similarity_score:.2f}")
    else:
        st.warning("Please upload a resume before analyzing the transcript.")
